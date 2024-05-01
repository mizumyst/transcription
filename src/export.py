import sys

import pandas as pd

import json
from docx import Document
import xml.etree.ElementTree as ET
from pathlib import Path


def write_docx(data, output_file):
    document = Document()

    metadata = data['metadata']
    speaker_string = ", ".join(f"[{speaker}]" for speaker in metadata['speakers'].split())

    time = metadata['time'] // 1000
    minutes, seconds = time // 60, time % 60
    hours, minutes = minutes // 60, minutes % 60

    time_string = ""
    if hours:
        time_string += f"{hours}時間"
    if minutes or hours:
        time_string += f"{minutes}分"
    if seconds or minutes or hours:
        time_string += f"{seconds}秒"

    document.add_paragraph(f"""{metadata['location']}, {metadata['date']}, {speaker_string}
録音時間 {time_string}
    """)

    text_paragraph = document.add_paragraph()
    document.add_page_break()
    gloss_paragraph = document.add_paragraph()

    last_speaker = None
    epochs_elapsed = 0
    text_running = None
    gloss_running = None

    for segment in data['transcribe']:
        speaker = segment['speaker']
        text = segment['text']
        try:
            gloss = segment['gloss']
        except KeyError:
            gloss = text
        if gloss is None:
            gloss = text
        

        if text_running and text_running not in "。？！、":
            text_paragraph.add_run(" ")
        if gloss_running and gloss_running not in "。？！、":
            gloss_paragraph.add_run(" ")

        if speaker == last_speaker:
            text_paragraph.add_run(text)
            gloss_paragraph.add_run(gloss)

        else:
            last_speaker = speaker

            text_paragraph.add_run(f"[{speaker}] {text}")
            gloss_paragraph.add_run(f"[{speaker}] {gloss}")
        
        text_running = text[-1]
        gloss_running = gloss[-1]
        
        if segment['timestamp'][1] > (epochs_elapsed + 1) * 300:
            epochs_elapsed += 1
            text_paragraph.add_run(f" ({epochs_elapsed * 5}:00) ")
            gloss_paragraph.add_run(f" ({epochs_elapsed * 5}:00) ")
        
    document.save(output_file)


def get_time_index(time):
    global times
    if time == -1:
        times = []
        return

    if time in times:
        return times.index(time) + 1

    times.append(time)
    return len(times)


def write_xml(data, output_file, media_path):
    root = ET.Element('ANNOTATION_DOCUMENT', {
        'AUTHOR':                           '',
        'DATE':                             '2023-12-11T07:52:21-05:00',
        'FORMAT':                           '3.0',
        'VERSION':                          '3.0',
        'xmlns:xsi':                        'http://www.w3.org/2001/XMLSchema-instance',
        'xsi:noNamespaceSchemaLocation':    'http://www.mpi.nl/tools/elan/EAFv3.0.xsd'
    })

    header = ET.SubElement(root, 'HEADER', {
        'MEDIA_FILE':   str(media_path),
        'TIME_UNITS':   'milliseconds'
    })

    time_order = ET.SubElement(root, 'TIME_ORDER')

    tiers = dict()

    speakers = list(set(segment['speaker'] for segment in data['transcribe']))  # TODO: use the pickle for this

    for speaker in speakers:
        speaker_no = speaker.split('/')[0]
        tiers[f'tx@{speaker}'] = ET.SubElement(root, 'TIER', {
            'LINGUISTIC_TYPE_REF':  'default-lt',
            'PARTICIPANT':          speaker,
            'TIER_ID':              f'tx@{speaker_no}',
        })
        
        tiers[f'gl@{speaker}'] = ET.SubElement(root, 'TIER', {
            'LINGUISTIC_TYPE_REF':  'translation',
            'PARTICIPANT':          speaker,
            'PARENT_REF':           f'tx@{speaker_no}',
            'TIER_ID':              f'gl@{speaker_no}',
        })
        
        tiers[f'note@{speaker}'] = ET.SubElement(root, 'TIER', {
            'LINGUISTIC_TYPE_REF':  'translation',
            'PARTICIPANT':          speaker,
            'PARENT_REF':           f'tx@{speaker_no}',
            'TIER_ID':              f'note@{speaker_no}',
        })

    tree = ET.ElementTree(root)

    lingustic_type = ET.SubElement(root, 'LINGUISTIC_TYPE', {
        'GRAPHIC_REFERENCES': "false",
        'LINGUISTIC_TYPE_ID': "default-lt",
        'TIME_ALIGNABLE': "true"
    })

    lingustic_type = ET.SubElement(root, 'LINGUISTIC_TYPE', {
        'CONSTRAINTS': "Symbolic_Association",
        'GRAPHIC_REFERENCES': "false",
        'LINGUISTIC_TYPE_ID': "translation",
        'TIME_ALIGNABLE': "false"
    })

    CONSTRAINTS = [
        ("Time_Subdivision", "Time subdivision of parent annotation's time interval, no time gaps allowed within this interval"),
        ("Symbolic_Subdivision", "Symbolic subdivision of a parent annotation. Annotations refering to the same parent are ordered"),
        ("Symbolic_Association", "1-1 association with a parent annotation"),
        ("Included_In", "Time alignable annotations within the parent annotation's time interval, gaps are allowed")
    ]

    for stereotype, description in CONSTRAINTS:
        ET.SubElement(root, 'CONSTRAINT', {
            'DESCRIPTION': description,
            'STEREOTYPE': stereotype
        })

    get_time_index(-1)

    annotation_count = 0
    for segment in data['transcribe']:
        start_index, end_index = segment['timestamp']
        start_index, end_index = get_time_index(start_index), get_time_index(end_index)
        
        speaker = segment['speaker']
        
        if 'text' in segment:
            annotation_count += 1
            tx_annotation_id = f"a{annotation_count}"
            
            container = ET.SubElement(tiers[f'tx@{speaker}'], 'ANNOTATION')
            annotation = ET.SubElement(container, 'ALIGNABLE_ANNOTATION', {
                'ANNOTATION_ID': tx_annotation_id,
                'TIME_SLOT_REF1': f"ts{start_index}",
                'TIME_SLOT_REF2': f"ts{end_index}",
            })
            value = ET.SubElement(annotation, 'ANNOTATION_VALUE')
            value.text = segment['text']

        if 'gloss' in segment:
            annotation_count += 1

            container = ET.SubElement(tiers[f'gl@{speaker}'], 'ANNOTATION')
            annotation = ET.SubElement(container, 'REF_ANNOTATION', {
                'ANNOTATION_ID': f"a{annotation_count}",
                'ANNOTATION_REF': tx_annotation_id,
            })
            value = ET.SubElement(annotation, 'ANNOTATION_VALUE')
            value.text = segment['gloss']

        if 'note' in segment:
            annotation_count += 1

            container = ET.SubElement(tiers[f'note@{speaker}'], 'ANNOTATION')
            annotation = ET.SubElement(container, 'REF_ANNOTATION', {
                'ANNOTATION_ID': f"a{annotation_count}",
                'ANNOTATION_REF': tx_annotation_id,
            })
            value = ET.SubElement(annotation, 'ANNOTATION_VALUE')
            value.text = segment['note']


    for index, time in enumerate(times, start=1):
        ET.SubElement(time_order, 'TIME_SLOT', {
            'TIME_SLOT_ID': f"ts{index}",
            'TIME_VALUE': str(int(time * 1000))
        })

    ET.indent(tree, space='\t', level=0)
    tree.write(output_file, encoding='utf-8')


def read_xml(input_file, metadata, translate):
    tree = ET.parse(input_file)
    root = tree.getroot()

    time_dict = {}
    times = root.findall('TIME_ORDER')[0]
    end_time = 0
    for time_slot in times:
        current_time_ms = int(time_slot.get('TIME_VALUE'))
        end_time = max(end_time, current_time_ms)
        time_dict[time_slot.get('TIME_SLOT_ID')] = current_time_ms / 1000

    tiers = root.findall('TIER')
    text, gloss, note = tiers[::3], tiers[1::3], tiers[2::3]

    text_refs = []
    text_dict = {}

    for text_tier, gloss_tier, note_tier in zip(text, gloss, note):
        speaker_id = text_tier.get('TIER_ID').split('@', maxsplit=1)[1]

        for annotation in text_tier:
            annotation = annotation[0]

            annotation_id = annotation.get('ANNOTATION_ID')
            start, end = (
                time_dict[annotation.get('TIME_SLOT_REF1')],
                time_dict[annotation.get('TIME_SLOT_REF2')]
            )

            text_refs.append(annotation_id)
            text_dict[annotation_id] = {
                'speaker': speaker_id,
                'text': annotation[0].text,
                'timestamp': [start, end]
                }

        for annotation in gloss_tier:
            annotation = annotation[0]
            text_dict[annotation.get('ANNOTATION_REF')]['gloss'] = annotation[0].text

        for annotation in note_tier:
            annotation = annotation[0]
            text_dict[annotation.get('ANNOTATION_REF')]['note'] = annotation[0].text 
        
    segments = list(text_dict.values())
    segments = sorted(segments, key=lambda annotation: annotation['timestamp'])

    metadata['time'] = end_time

    return {
        'metadata': metadata,
        'transcribe': segments,
        'translate': translate
    }


if __name__ == '__main__':
    directory, target, *other = sys.argv    # TODO: one day this will be done with argparse
    source = None if not other else other[0]

    directory = Path(directory)
    metadata = pd.read_csv(directory / 'manifest.csv').set_index('filename').to_dict(orient='index')

    match target:
        case 'xml':
            for path in Path(directory).glob("*.json"):
                with open(path, 'r', encoding='utf-8') as file:
                    data = json.load(file)

                    write_xml(data, path.with_suffix('.eaf'), path.with_suffix(''))

        case 'docx':
            if source == 'xml':
                for path in directory.glob("*.eaf"):
                    with open(path, 'r', encoding='utf-8') as file:
                        data = read_xml(file, metadata[path.stem], "")
                        write_docx(data, path.with_suffix('.docx'))
            
            elif source == 'json':
                for path in directory.glob("*.json"):
                    with open(path, 'r', encoding='utf-8') as file:
                        data = json.load(file)
                        write_docx(data, path.with_suffix('.docx'))
                
        case 'json':
            for path in directory.glob("*.eaf"):
                with open(path, 'r', encoding='utf-8') as file:
                    data = read_xml(file, metadata[path.stem], "")
            
                with open(path.with_suffix('.json'), 'w', encoding='utf-8') as file:
                    json.dump(data, file)
